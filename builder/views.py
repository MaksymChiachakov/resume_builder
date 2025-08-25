from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
import io
from docx import Document
from .models import ResumeTemplate, Resume, Announcement, ResumeSection
from .forms import ResumeForm, UserUpdateForm, AnnouncementForm

def home(request):
    templates = ResumeTemplate.objects.all()[:3]
    announcements = Announcement.objects.order_by('-created_at')[:5]
    return render(request, 'builder/home.html', {
        'templates': templates,
        'announcements': announcements
    })

def resume_templates(request):
    templates = ResumeTemplate.objects.all()
    return render(request, 'builder/templates.html', {'templates': templates})

@login_required
def resume_list(request):
    resumes = Resume.objects.filter(user=request.user)
    return render(request, 'builder/resume_list.html', {'resumes': resumes})

@login_required
def profile(request):
    if request.method == 'POST':
        form = UserUpdateForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Профіль успішно оновлено!')
            return redirect('profile')
    else:
        form = UserUpdateForm(instance=request.user)
    
    resumes = Resume.objects.filter(user=request.user)
    return render(request, 'builder/profile.html', {
        'form': form,
        'resumes': resumes
    })

def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password1')
            user = authenticate(username=username, password=password)
            login(request, user)
            messages.success(request, f'Вітаємо, {username}! Ваш акаунт успішно створено.')
            return redirect('home')
    else:
        form = UserCreationForm()
    return render(request, 'builder/register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            messages.success(request, f'Вітаємо, {user.username}! Ви увійшли в систему.')
            return redirect('home')
    else:
        form = AuthenticationForm()
    return render(request, 'builder/login.html', {'form': form})

def logout_view(request):
    logout(request)
    messages.info(request, 'Ви вийшли з системи.')
    return redirect('home')

@login_required
def create_resume(request):
    if request.method == 'POST':
        form = ResumeForm(request.POST, request.FILES)
        if form.is_valid():
            resume = form.save(commit=False)
            resume.user = request.user
            resume.save()
            
            # Додаємо стандартні секції
            default_sections = [
                ('Особиста інформація', 'Введіть вашу контактну інформацію'),
                ('Досвід роботи', 'Опишіть ваш досвід роботи'),
                ('Освіта', 'Додайте інформацію про освіту'),
                ('Навички', 'Перелічіть ваші навички'),
            ]
            
            for i, (title, content) in enumerate(default_sections):
                ResumeSection.objects.create(
                    resume=resume,
                    title=title,
                    content=content,
                    order=i
                )
                
            messages.success(request, 'Резюме успішно створено! Тепер ви можете додати інформацію.')
            return redirect('edit_resume_sections', resume_id=resume.id)
    else:
        form = ResumeForm()
    return render(request, 'builder/create_resume.html', {'form': form})

@login_required
def edit_resume_sections(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    
    if request.method == 'POST':
        # Оновлення існуючих секцій
        if 'save_sections' in request.POST:
            for section in resume.sections.all():
                content_field_name = f'section_content_{section.id}'
                if content_field_name in request.POST:
                    section.content = request.POST[content_field_name]
                    section.save()
            messages.success(request, 'Зміни збережено!')
            return redirect('edit_resume_sections', resume_id=resume.id)
        
        # Додавання нової секції
        elif 'add_section' in request.POST:
            section_title = request.POST.get('section_title')
            section_content = request.POST.get('section_content')
            
            if section_title and section_content:
                last_order = ResumeSection.objects.filter(resume=resume).order_by('-order').first()
                new_order = last_order.order + 1 if last_order else 0
                
                ResumeSection.objects.create(
                    resume=resume,
                    title=section_title,
                    content=section_content,
                    order=new_order
                )
                messages.success(request, 'Розділ додано!')
                return redirect('edit_resume_sections', resume_id=resume.id)
    
    sections = ResumeSection.objects.filter(resume=resume).order_by('order')
    return render(request, 'builder/edit_resume.html', {
        'resume': resume,
        'sections': sections
    })

@login_required
def delete_resume_section(request, section_id):
    section = get_object_or_404(ResumeSection, id=section_id)
    if section.resume.user != request.user:
        messages.error(request, 'У вас немає прав для цієї дії.')
        return redirect('resume_list')
    
    resume_id = section.resume.id
    section.delete()
    messages.success(request, 'Розділ видалено!')
    return redirect('edit_resume_sections', resume_id=resume_id)

@login_required
def delete_resume(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    
    if request.method == 'POST':
        resume.delete()
        messages.success(request, 'Резюме видалено!')
        return redirect('resume_list')
    
    return render(request, 'builder/delete_resume.html', {'resume': resume})

@login_required
def clone_resume(request, resume_id):
    original_resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    
    # Створюємо копію резюме
    cloned_resume = Resume.objects.create(
        user=request.user,
        template=original_resume.template,
        title=f"Копія {original_resume.title}",
        photo=original_resume.photo
    )
    
    # Копіюємо всі секції
    for section in original_resume.sections.all():
        ResumeSection.objects.create(
            resume=cloned_resume,
            title=section.title,
            content=section.content,
            order=section.order
        )
    
    messages.success(request, 'Резюме успішно скопійовано!')
    return redirect('edit_resume_sections', resume_id=cloned_resume.id)

def announcements(request):
    news = Announcement.objects.order_by('-created_at')
    return render(request, 'builder/announcements.html', {'news': news})

def is_admin(user):
    return user.is_staff

@user_passes_test(is_admin)
def create_announcement(request):
    if request.method == 'POST':
        form = AnnouncementForm(request.POST)
        if form.is_valid():
            announcement = form.save(commit=False)
            announcement.author = request.user
            announcement.save()
            messages.success(request, 'Оголошення створено!')
            return redirect('announcements')
    else:
        form = AnnouncementForm()
    return render(request, 'builder/create_announcement.html', {'form': form})

@user_passes_test(is_admin)
def edit_announcement(request, announcement_id):
    announcement = get_object_or_404(Announcement, id=announcement_id)
    if request.method == 'POST':
        form = AnnouncementForm(request.POST, instance=announcement)
        if form.is_valid():
            form.save()
            messages.success(request, 'Оголошення оновлено!')
            return redirect('announcements')
    else:
        form = AnnouncementForm(instance=announcement)
    return render(request, 'builder/edit_announcement.html', {'form': form, 'announcement': announcement})

@user_passes_test(is_admin)
def delete_announcement(request, announcement_id):
    announcement = get_object_or_404(Announcement, id=announcement_id)
    if request.method == 'POST':
        announcement.delete()
        messages.success(request, 'Оголошення видалено!')
        return redirect('announcements')
    return render(request, 'builder/delete_announcement.html', {'announcement': announcement})

@login_required
def view_resume(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    sections = ResumeSection.objects.filter(resume=resume).order_by('order')
    return render(request, 'builder/view_resume.html', {
        'resume': resume,
        'sections': sections
    })

@login_required
def export_pdf(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    sections = ResumeSection.objects.filter(resume=resume).order_by('order')
    
    template_path = 'builder/export_pdf.html'
    context = {'resume': resume, 'sections': sections}
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{resume.title}.pdf"'
    
    template = get_template(template_path)
    html = template.render(context)
    
    pisa_status = pisa.CreatePDF(html, dest=response)
    
    if pisa_status.err:
        return HttpResponse('Помилка при створенні PDF')
    return response

@login_required
def export_docx(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    sections = ResumeSection.objects.filter(resume=resume).order_by('order')
    
    document = Document()
    document.add_heading(resume.title, 0)
    
    for section in sections:
        document.add_heading(section.title, level=1)
        document.add_paragraph(section.content)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{resume.title}.docx"'
    
    document.save(response)
    return response